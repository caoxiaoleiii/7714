import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_font(run, latin="Calibri", east_asia="宋体", size=10.5, bold=False, color=None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def apply_styles(doc):
    section = doc.sections[-1]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.16

    heading = doc.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading.font.size = Pt(15)
    heading.font.color.rgb = RGBColor(0, 0, 0)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(8)


def clear_headers_footers(doc):
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                paragraph.text = ""


def add_section(doc, heading, items, first=False):
    if not first:
        doc.add_section(WD_SECTION.NEW_PAGE)
        apply_styles(doc)
    h = doc.add_paragraph(style="Heading 1")
    run = h.add_run(heading)
    set_font(run, east_asia="黑体", size=15, bold=True)

    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.16
        set_font(p.add_run(item))


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: python make_reference_docx.py input_sections.json output.docx")

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    data = json.loads(input_path.read_text(encoding="utf-8"))
    sections = data.get("sections") or []
    if not sections:
        raise SystemExit("input JSON must contain a non-empty 'sections' array")

    doc = Document()
    apply_styles(doc)
    for idx, section in enumerate(sections):
        heading = section.get("heading") or f"第{idx + 1}章参考文献"
        items = section.get("items") or []
        add_section(doc, heading, items, first=(idx == 0))

    clear_headers_footers(doc)
    doc.save(output_path)


if __name__ == "__main__":
    main()
